import requests
from bs4 import BeautifulSoup
from docx import Document
from tqdm import tqdm
import argparse

def fetch_html(url):
    """
    Fetches the HTML content from the given URL.
    
    :param url: The URL to fetch the content from.
    :return: The HTML content as a string if successful, otherwise None.
    """
    response = requests.get(url)
    response.encoding = 'utf-8'  # Ensure correct encoding
    
    if response.status_code == 200:
        return response.text
    else:
        print(f"Failed to retrieve the page: {url}. Status code: {response.status_code}")
        return None

def parse_sanskrit_content(html_content):
    """
    Parses the Sanskrit content from the provided HTML content.
    
    :param html_content: The HTML content to parse.
    :return: A list of Sanskrit content strings.
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    padapatha_divs = soup.find_all('div', class_='pada_dev_acc')
    
    sanskrit_texts = []
    for div in padapatha_divs:
        texts = div.find_all('span', class_='sanskrit')
        sanskrit_texts.append(' '.join([text.get_text(strip=True) for text in texts]))
    
    return sanskrit_texts

def add_content_to_document(doc, sanskrit_texts, y_value):
    """
    Adds the Sanskrit content to a Word document.
    
    :param doc: The Word document object.
    :param sanskrit_texts: The list of Sanskrit content strings.
    :param y_value: The value of 'y' to label the sections.
    """
    for i, text in enumerate(sanskrit_texts, start=1):
        # doc.add_heading(f'Section {i} (y={y_value})', level=2)
        doc.add_paragraph(text)
        # doc.add_paragraph("\n")  # Add some space between sections

def process_urls(x, y1, y2, base_url, doc):
    """
    Processes the URLs in the specified range and adds the content to the Word document.
    
    :param x: The first part of the URL.
    :param y1: The starting number of the second part of the URL.
    :param y2: The ending number of the second part of the URL.
    :param base_url: The base URL to be used.
    :param doc: The Word document object.
    """
    for y in tqdm(range(y1, y2 + 1), desc="Processing URLs"):
        y_formatted = f'{y:03d}'
        url = f'{base_url}{x}/{x}-{y_formatted}.htm'
        
        html_content = fetch_html(url)
        if html_content:
            sanskrit_texts = parse_sanskrit_content(html_content)
            add_content_to_document(doc, sanskrit_texts, y_formatted)
        else:
            print("doc search ended")
            break

def main():
    # Set up argument parser
    parser = argparse.ArgumentParser(description="Extract Sanskrit content from Rigveda HTML pages and save to a Word document.")
    parser.add_argument('x', type=str, help="The first part of the URL (e.g., '01' to '10').")
    parser.add_argument('y1', type=int, help="The starting number of the second part of the URL (e.g., 1).")
    parser.add_argument('y2', type=int, help="The ending number of the second part of the URL (e.g., 100).")
    parser.add_argument('output', type=str, help="The output path for .docx file.",  default="output.docx", nargs='?')
    
    args = parser.parse_args()
    
    # Base URL
    base_url = 'https://sri-aurobindo.co.in/workings/matherials/rigveda/'
    
    # Create a new Word document
    doc = Document()
    
    # Process the URLs
    process_urls(args.x, args.y1, args.y2, base_url, doc)
    
    # Save the document to a file
    output_path = args.output
    doc.save(output_path)
    print(f"Sanskrit content saved to {output_path}")

if __name__ == "__main__":
    main()
